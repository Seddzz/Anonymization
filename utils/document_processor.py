"""
Document Processor - Handles anonymization while preserving document structure
"""
import pdfplumber
import docx
from docx import Document
import os
import tempfile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

class DocumentProcessor:
    def __init__(self, pipeline):
        """
        Initialize with an anonymization pipeline
        """
        self.pipeline = pipeline
    
    def anonymize_txt(self, file_path, output_path=None):
        """
        Anonymize a TXT file while preserving structure
        """
        try:
            # Read text content
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Anonymize the text
            anonymized_text = self.pipeline.anonymize(text_content)
            
            # Get replacement mapping for statistics
            replacement_details = self.pipeline.replacer.get_replacements_with_types()
            replacement_mapping = {}
            for original, replacement, entity_type in replacement_details:
                replacement_mapping[original] = replacement
            
            # Create output file
            if output_path is None:
                output_path = file_path.replace('.txt', '_anonymized.txt')
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(anonymized_text)
            
            return {
                'success': True,
                'output_path': output_path,
                'original_text': text_content,
                'anonymized_text': anonymized_text,
                'replacement_mapping': replacement_mapping,
                'message': f'TXT file anonymized successfully: {os.path.basename(output_path)}',
                'file_type': 'txt'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing TXT: {str(e)}'
            }

    def anonymize_pdf(self, file_path, output_path=None):
        """
        Anonymize a PDF file while creating a proper PDF output
        """
        try:
            # Extract text from PDF
            text_content = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n\n"
            
            # Anonymize the text
            anonymized_text = self.pipeline.anonymize(text_content)
            
            # Get replacement mapping for statistics
            replacement_details = self.pipeline.replacer.get_replacements_with_types()
            replacement_mapping = {}
            for original, replacement, entity_type in replacement_details:
                replacement_mapping[original] = replacement
            
            # Create new PDF with anonymized content
            if output_path is None:
                output_path = file_path.replace('.pdf', '_anonymized.pdf')
            
            self._create_pdf_from_text(anonymized_text, output_path)
            
            return {
                'success': True,
                'output_path': output_path,
                'original_text': text_content,
                'anonymized_text': anonymized_text,
                'replacement_mapping': replacement_mapping,
                'message': f'PDF anonymized successfully: {os.path.basename(output_path)}',
                'file_type': 'pdf'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing PDF: {str(e)}'
            }
    
    def anonymize_docx(self, file_path, output_path=None):
        """
        Anonymize a DOCX file while preserving structure and formatting
        """
        try:
            # Load the document
            doc = Document(file_path)
            
            # Extract all text for anonymization mapping
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            # Get anonymization mappings
            entities = self.pipeline.detector.detect(full_text)
            
            # Get replacement mapping from replacer
            anonymized_full_text = self.pipeline.anonymize(full_text)
            replacer_details = self.pipeline.replacer.get_replacements_with_types()
            
            replacement_mapping = {}
            for original, replacement, entity_type in replacer_details:
                replacement_mapping[original] = replacement
            
            # Apply replacements to each paragraph while preserving formatting
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    original_text = paragraph.text
                    new_text = original_text
                    
                    # Apply all replacements
                    for original, replacement in replacement_mapping.items():
                        new_text = new_text.replace(original, replacement)
                    
                    # Update paragraph text if it changed
                    if new_text != original_text:
                        # Clear existing runs and add new text
                        paragraph.clear()
                        paragraph.add_run(new_text)
            
            # Save anonymized document
            if output_path is None:
                output_path = file_path.replace('.docx', '_anonymized.docx')
            
            doc.save(output_path)
            
            return {
                'success': True,
                'output_path': output_path,
                'original_text': full_text,
                'anonymized_text': anonymized_full_text,
                'replacement_mapping': replacement_mapping,
                'message': f'DOCX anonymized successfully: {os.path.basename(output_path)}',
                'file_type': 'docx'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing DOCX: {str(e)}'
            }

    def process_file(self, file_path, file_type=None):
        """
        Process any supported file type
        """
        if file_type is None:
            file_type = file_path.split('.')[-1].lower()
        
        if file_type == 'pdf':
            return self.anonymize_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return self.anonymize_docx(file_path)
        elif file_type == 'txt':
            return self.anonymize_txt(file_path)
        else:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_type}'
            }

    def _create_pdf_from_text(self, text, output_path):
        """
        Create a professional PDF from text content using ReportLab
        """
        try:
            # Create PDF document with proper styling
            doc = SimpleDocTemplate(
                output_path, 
                pagesize=A4,
                rightMargin=72, 
                leftMargin=72,
                topMargin=72, 
                bottomMargin=18
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Split text into paragraphs
            paragraphs = text.split('\n\n')
            
            # Create story (content) for PDF
            story = []
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Clean up the text
                    cleaned_text = para_text.strip().replace('\n', ' ')
                    
                    # Create paragraph with proper styling
                    para = Paragraph(cleaned_text, styles['Normal'])
                    story.append(para)
                    story.append(Spacer(1, 12))  # Add space between paragraphs
            
            # Build PDF
            doc.build(story)
            
        except Exception as e:
            print(f"Error with advanced PDF creation: {e}")
            # Fallback to simple PDF creation
            self._create_simple_pdf(text, output_path)
    
    def _create_simple_pdf(self, text, output_path):
        """
        Fallback method to create a simple PDF using canvas
        """
        try:
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            
            # Set up text formatting
            c.setFont("Helvetica", 11)
            
            # Split text into lines
            lines = text.replace('\n\n', '\n').split('\n')
            y_position = height - 50  # Start near top of page
            line_height = 14
            
            for line in lines:
                if y_position < 50:  # Start new page if needed
                    c.showPage()
                    y_position = height - 50
                    c.setFont("Helvetica", 11)
                
                # Handle long lines by wrapping
                if len(line) > 85:  # Wrap at ~85 characters
                    words = line.split(' ')
                    current_line = ""
                    
                    for word in words:
                        if len(current_line + word) < 85:
                            current_line += word + " "
                        else:
                            if current_line:
                                c.drawString(50, y_position, current_line.strip())
                                y_position -= line_height
                                if y_position < 50:
                                    c.showPage()
                                    y_position = height - 50
                                    c.setFont("Helvetica", 11)
                            current_line = word + " "
                    
                    if current_line:
                        c.drawString(50, y_position, current_line.strip())
                        y_position -= line_height
                else:
                    c.drawString(50, y_position, line)
                    y_position -= line_height
            
            c.save()
            
        except Exception as e:
            print(f"Error with simple PDF creation: {e}")
            raise
